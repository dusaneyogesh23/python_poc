#!/usr/bin/python
# - * - coding: UTF-8 - * -
import sys
import docx
import MySQLdb as master_db
from docx import Document

UNKNOW_DATABASE_ERR = 1049
MASTER_DATABASE = "master_data_test5"

create_master_tbl = """ CREATE TABLE IF NOT EXISTS master_tbl (
			       ID integer PRIMARY KEY AUTO_INCREMENT,
                               Attr_Number text,
                               Attr_Name text,
                               Attr_Discr text,
                               Attr_Type text,
                               Attr_Size text,
                               Attr_Value text,
                               Attr_Access_Super_Mode text,
                               Attr_Access_USer_Mode text,
                               Attr_Excep_Note text,
                               Attr_Customer_Face text
                               );"""

unique_key_master_tbl = "ALTER TABLE master_tbl ADD UNIQUE Attr_Number(Attr_Number(6));"

unique_key = "alter ignore table master_tbl add unique index dwarf_Attr_Number (Attr_Number(5));"

insert_query_master_tbl = """ INSERT INTO master_tbl (Attr_Number,
                                    Attr_Name, 
                                    Attr_Discr,
                                    Attr_Type,
                                    Attr_Size,
                                    Attr_Value,
                                    Attr_Access_Super_Mode,
                                    Attr_Access_USer_Mode,
                                    Attr_Excep_Note,
                                    Attr_Customer_Face
                                 ) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"""

column_count = """ SELECT COUNT(*) AS Columns FROM INFORMATION_SCHEMA.COLUMNS 
		   WHERE table_schema = 'master_data_test5' 
		   AND table_name = 'master_tbl' """

def create_database(master_db_name):
	
	try:
		conn = master_db.connect("localhost","root","root")
		cursor = conn.cursor()
		master_db_query = "CREATE DATABASE IF NOT EXISTS " + master_db_name
		cursor.execute("SET sql_notes = 0; ")
		cursor.execute(master_db_query)
		cursor.execute("SET sql_notes = 1; ")
	
		conn = master_db.connect("localhost","root","root",master_db_name,use_unicode=True, charset="utf8")
		return conn
	except master_db.Error as err:
                print(err)

	return None

def create_database_connection(master_db_name):
	
	try:	
		# Open database connection
                conn = master_db.connect("localhost","root","root",master_db_name,use_unicode=True, charset="utf8")
		return conn

	except master_db.Error as err:

		print(err)
		if err.args[0] == UNKNOW_DATABASE_ERR :
			conn = create_database(master_db_name)
			return conn
	return None

def create_master_table(conn,create_tbl):
	try:
		cursor = conn.cursor()
		cursor.execute("SET sql_notes = 0; ")
		cursor.execute(create_tbl)
		#cursor.execute(unique_key)
		cursor.execute("SET sql_notes = 1; ")
	except master_db.Error as err:
                print(err)

def db_insert_single_records(conn,row_data):
	try:
		cursor = conn.cursor()
        	cursor.execute(insert_query_master_tbl,row_data)
	except master_db.Error as err:
		print(err)

def db_insert_multiple_records(conn,row_data):
        try:
                cursor = conn.cursor()
                cursor.executemany(insert_query_master_tbl,row_data)
        except master_db.Error as err:
                print(err)

def db_save_records(conn,record_count):
	try:
		conn.commit()
		conn.close()
		print ( str(record_count) +" - Records saved successfully !!!\r\n");
	except master_db.Error as err:
                print(err)

def filter_non_acsii_char(text):
        text = text.replace(unicode('“',"utf-8"),"").replace(unicode('”',"utf-8"),"")
        text = text.replace(unicode('‘',"utf-8"),"").replace(unicode('’',"utf-8"),"")
        text = text.strip("'").strip('"')
	text = text.replace(" ","")
	text = "'" + text + "'"
        return text

def open_read_docx(conn,doc_name):

	print("Start Reading the Master Word Document !!!")

	try : 
		document = Document()
   		document = docx.Document(doc_name)
	except :
		print("Error in reading "+ doc_name +" Document !!! \r\nPlease check file format and it's path !!!\r\n")
		return 0
		
	master_list = list()
	total_record = 0
	
	for table in document.tables:
		first_cell = table.rows[1].cells[0].text
		if not (first_cell == "Attribute Number"):
			continue
		
		col_count = len(table.columns)

		#if not (col_count == 10):
		#	print("Currputed word Document or Invalide document")
		#	return 0
		
		for i, row in enumerate(table.rows):
			if i <= 1:
				continue

			attrbute_number = row.cells[0].text
			attrbute_number = attrbute_number.replace(" ","")
			if not attrbute_number.isdigit() :
				continue

			row_list = list()

			total_record += 1

			for cell in row.cells:
				text = cell.text 
				#text = filter_non_acsii_char(cell.text)
				print text
				row_list.append(text) 
				#sys.stdout.write("-")
			#print row_list
			#db_insert_single_records(conn,row_list)
			#master_list.append(row_list)
	#print master_list
	db_insert_multiple_records(conn,master_list)
	return total_record

def main():
	if len (sys.argv) != 2 :
		print ("Error : Please enter valid Document path and name (Ex. rsm.docx) !!!")
		print ("Usage : Ex. "+sys.argv[0] + " rsm.docx")
		sys.exit (1)

	doc_name = sys.argv[1]
	if doc_name.endswith('.docx'):
		print("Master Database Generator script !!!")

		conn = create_database_connection( MASTER_DATABASE )
		create_master_table( conn, create_master_tbl )

		record_count = open_read_docx(conn,doc_name)

		if record_count > 0 :
			db_save_records(conn,record_count)
	else:
		print("Please enter valid MS Word document format file !!!")

if __name__ == '__main__':
   main()

